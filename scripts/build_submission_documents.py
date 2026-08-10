"""Build and style the final Word/PDF submission documents.

Run with the Codex primary runtime because it owns the document dependencies.
Pandoc is also required so LaTeX expressions become native Word Office Math:

    "$CODEX_PRIMARY_RUNTIME_PYTHON" scripts/build_submission_documents.py
"""

from __future__ import annotations

import io
import re
import shutil
import subprocess
import zipfile
from functools import cache
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"

NAVY = "18324A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
CYAN = "2F7F9D"
INK = "243340"
MUTED = "5D6975"
LIGHT = "F2F4F7"
PALE_BLUE = "EAF2F7"
WHITE = "FFFFFF"
RULE = "CFD8E1"


def _rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def _set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = _rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, *, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_table_geometry(table, widths_dxa: list[int], *, indent_dxa: int = 120) -> None:
    if sum(widths_dxa) != 9360:
        raise ValueError(f"Table widths must total 9360 DXA; got {sum(widths_dxa)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _column_widths(rows: list[list[str]]) -> list[int]:
    columns = len(rows[0])
    weights: list[float] = []
    for index in range(columns):
        values = [row[index] for row in rows]
        longest = max(len(value) for value in values)
        numeric = all(
            bool(re.fullmatch(r"[+−\-0-9.%$=\s]+", value.strip()))
            for value in values[1:]
            if value.strip()
        )
        weight = min(max(longest, 7), 30)
        if numeric:
            weight = min(weight, 12)
        weights.append(float(weight))
    raw = [round(9360 * weight / sum(weights)) for weight in weights]
    raw[-1] += 9360 - sum(raw)
    return raw


def _add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    _set_run_font(run, size=8.5, color=MUTED)
    for field_name in ("PAGE", "NUMPAGES"):
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = field_name
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        value = OxmlElement("w:t")
        value.text = "1"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        field_run = paragraph.add_run()
        field_run._r.extend([begin, instruction, separate, value, end])
        _set_run_font(field_run, size=8.5, color=MUTED)
        if field_name == "PAGE":
            slash = paragraph.add_run(" / ")
            _set_run_font(slash, size=8.5, color=MUTED)


def _configure_section(section, *, label: str, subtitle: str, first_page: bool = False) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = first_page

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    left = paragraph.add_run(label)
    _set_run_font(left, size=8.5, color=MUTED, bold=True)
    right = paragraph.add_run(f"    {subtitle}")
    _set_run_font(right, size=8.5, color=MUTED)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    _add_page_field(footer_paragraph)


def _set_style(
    style,
    *,
    font_size: float,
    color: str,
    bold: bool = False,
    italic: bool = False,
    before: float = 0,
    after: float = 0,
    line: float = 1.0,
    keep_with_next: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(font_size)
    style.font.color.rgb = _rgb(color)
    style.font.bold = bold
    style.font.italic = italic
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_with_next
    fmt.alignment = alignment


def _setup_styles(doc: Document, *, compact: bool) -> None:
    styles = doc.styles
    body_size = 10 if compact else 11
    body_after = 4 if compact else 6
    body_line = 1.05 if compact else 1.10
    _set_style(
        styles["Normal"],
        font_size=body_size,
        color=INK,
        after=body_after,
        line=body_line,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY if compact else WD_ALIGN_PARAGRAPH.LEFT,
    )
    _set_style(
        styles["Heading 1"],
        font_size=16,
        color=BLUE,
        bold=True,
        before=16,
        after=8,
        line=1.0,
        keep_with_next=True,
    )
    _set_style(
        styles["Heading 2"],
        font_size=13,
        color=BLUE,
        bold=True,
        before=12,
        after=6,
        line=1.0,
        keep_with_next=True,
    )
    _set_style(
        styles["Heading 3"],
        font_size=12,
        color=DARK_BLUE,
        bold=True,
        before=8,
        after=4,
        line=1.0,
        keep_with_next=True,
    )
    _set_style(
        styles["Caption"],
        font_size=8.5,
        color=MUTED,
        italic=True,
        before=4,
        after=6,
        line=1.0,
        keep_with_next=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    if "Equation" not in styles:
        equation = styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
    else:
        equation = styles["Equation"]
    _set_style(
        equation,
        font_size=10.5,
        color=NAVY,
        italic=True,
        before=5,
        after=7,
        line=1.0,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Callout"]
    _set_style(
        callout,
        font_size=10.5 if compact else 11,
        color=NAVY,
        bold=True,
        before=5,
        after=7,
        line=1.08,
    )


def _add_numbering(doc: Document, *, ordered: bool, compact: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.extend([start, num_fmt, lvl_text, suffix])
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "40" if compact else "80")
    spacing.set(qn("w:line"), "260" if compact else "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, indent, spacing])
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def _apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    p_pr.append(num_pr)


INLINE = re.compile(r"(\$[^$\n]+\$|\*\*.*?\*\*|`.*?`|\*.*?\*|<https?://[^>]+>)")


@cache
def _latex_to_omml_xml(latex: str) -> bytes:
    """Convert one LaTeX expression to native Word Office Math XML.

    Pandoc's TeX reader and DOCX writer provide a deterministic, standards-based
    conversion to OMML.  Keeping the cached representation as XML bytes ensures
    that every insertion receives a fresh element with no shared parent.
    """

    expression = latex.strip()
    if not expression:
        raise ValueError("Cannot render an empty mathematical expression")

    pandoc = shutil.which("pandoc")
    if pandoc is None:
        raise RuntimeError(
            "Pandoc is required to render report equations as native Office Math. "
            "Install pandoc and rerun scripts/build_submission_documents.py."
        )

    completed = subprocess.run(
        [
            pandoc,
            "--from=markdown+tex_math_dollars",
            "--to=docx",
            "--output=-",
        ],
        input=f"${expression}$\n".encode(),
        capture_output=True,
        check=False,
    )
    if completed.returncode != 0:
        message = completed.stderr.decode("utf-8", errors="replace").strip()
        raise RuntimeError(f"Pandoc could not render equation {expression!r}: {message}")

    try:
        with zipfile.ZipFile(io.BytesIO(completed.stdout)) as archive:
            document_xml = archive.read("word/document.xml")
    except (KeyError, zipfile.BadZipFile) as exc:
        raise RuntimeError(f"Pandoc returned an invalid DOCX for equation {expression!r}") from exc

    document = parse_xml(document_xml)
    equation = next(document.iter(qn("m:oMath")), None)
    if equation is None:
        raise RuntimeError(f"Pandoc returned no Office Math object for equation {expression!r}")
    return etree.tostring(equation, encoding="utf-8")


def _add_math(paragraph, latex: str) -> None:
    paragraph._p.append(parse_xml(_latex_to_omml_xml(latex)))


def _add_inline(paragraph, text: str, *, size: float | None = None, color: str = INK) -> None:
    cursor = 0
    for match in INLINE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            _set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("$"):
            _add_math(paragraph, token[1:-1])
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            _set_run_font(run, size=size, color=color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, name="Consolas", size=(size or 10) - 0.3, color=DARK_BLUE)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, size=size, color=color, italic=True)
        else:
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, size=size, color=BLUE)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        _set_run_font(run, size=size, color=color)


def _add_table(doc: Document, rows: list[list[str]], *, compact: bool) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    widths = _column_widths(rows)
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            is_numeric = bool(re.fullmatch(r"[+−\-0-9.%$=\s]+", value.strip()))
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if is_numeric or row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            )
            _add_inline(
                paragraph,
                value,
                size=7.6 if compact and len(rows[0]) >= 5 else 8.3 if compact else 9,
                color=WHITE if row_index == 0 else INK,
            )
            if row_index == 0:
                _set_cell_shading(cell, NAVY)
                for run in paragraph.runs:
                    run.bold = True
            elif row_index % 2 == 0:
                _set_cell_shading(cell, LIGHT)
    _set_repeat_table_header(table.rows[0])
    _set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)


def _add_figure(
    doc: Document,
    alt: str,
    path: Path,
    *,
    compact: bool,
    width_inches: float | None = None,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run()
    width = width_inches if width_inches is not None else 5.65 if compact else 5.85
    shape = run.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", alt)
    caption = doc.add_paragraph(style="Caption")
    _add_inline(caption, alt, size=8.5, color=MUTED)


def _parse_markdown(
    doc: Document,
    path: Path,
    *,
    compact: bool,
    heading_shift: int = 0,
    skip_title_block: bool = True,
    front_matter_until: str | None = None,
    forced_break_before: set[str] | None = None,
    figure_width_inches: float | None = None,
) -> None:
    forced_break_before = forced_break_before or set()
    lines = path.read_text(encoding="utf-8").splitlines()
    index = 0
    current_num: int | None = None
    current_bullet: int | None = None
    in_equation = False
    equation_lines: list[str] = []
    front_matter = skip_title_block

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()

        if in_equation:
            if stripped == "$$":
                paragraph = doc.add_paragraph(style="Equation")
                _add_math(paragraph, " ".join(equation_lines))
                equation_lines.clear()
                in_equation = False
            else:
                equation_lines.append(stripped)
            index += 1
            continue
        if stripped == "$$":
            in_equation = True
            index += 1
            continue
        if not stripped:
            current_num = None
            current_bullet = None
            index += 1
            continue

        if front_matter_until and front_matter:
            if stripped == f"## {front_matter_until}":
                front_matter = False
            else:
                index += 1
                continue
        elif front_matter and stripped.startswith(("# ", "**")):
            index += 1
            continue
        else:
            front_matter = False

        image = re.fullmatch(r"!\[([^]]+)]\(([^)]+)\)", stripped)
        if image:
            image_path = (path.parent / image.group(2)).resolve()
            _add_figure(
                doc,
                image.group(1),
                image_path,
                compact=compact,
                width_inches=figure_width_inches,
            )
            index += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_rows: list[list[str]] = []
            while index < len(lines):
                candidate = lines[index].strip()
                if not (candidate.startswith("|") and candidate.endswith("|")):
                    break
                cells = [value.strip() for value in candidate.strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", value) for value in cells):
                    table_rows.append(cells)
                index += 1
            _add_table(doc, table_rows, compact=compact)
            continue

        heading = re.match(r"^(#{2,4})\s+(.*)$", stripped)
        if heading:
            title = heading.group(2)
            if title in forced_break_before:
                page = doc.add_paragraph()
                page.add_run().add_break(WD_BREAK.PAGE)
            level = max(1, min(3, len(heading.group(1)) - 1 + heading_shift))
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            _add_inline(paragraph, title, color=BLUE if level < 3 else DARK_BLUE)
            index += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            item_text = numbered.group(2)
            index += 1
            while index < len(lines) and lines[index].startswith(("  ", "\t")) and lines[index].strip():
                item_text += " " + lines[index].strip()
                index += 1
            if current_num is None:
                current_num = _add_numbering(doc, ordered=True, compact=compact)
            paragraph = doc.add_paragraph()
            _apply_num(paragraph, current_num)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.space_after = Pt(2 if compact else 4)
            _add_inline(paragraph, item_text)
            continue

        bullet = re.match(r"^-\s+(.*)$", stripped)
        if bullet:
            item_text = bullet.group(1)
            index += 1
            while index < len(lines) and lines[index].startswith(("  ", "\t")) and lines[index].strip():
                item_text += " " + lines[index].strip()
                index += 1
            if current_bullet is None:
                current_bullet = _add_numbering(doc, ordered=False, compact=compact)
            paragraph = doc.add_paragraph()
            _apply_num(paragraph, current_bullet)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.space_after = Pt(2 if compact else 4)
            _add_inline(paragraph, item_text)
            continue

        if stripped.startswith("> "):
            paragraph = doc.add_paragraph(style="Callout")
            _add_inline(paragraph, stripped[2:], color=NAVY)
            p_pr = paragraph._p.get_or_add_pPr()
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), PALE_BLUE)
            p_pr.append(shading)
            index += 1
            continue

        combined = stripped
        index += 1
        while index < len(lines):
            nxt = lines[index].strip()
            if not nxt or nxt.startswith(("#", "|", "![", "> ", "- ", "$$")) or re.match(r"^\d+\.\s", nxt):
                break
            combined += " " + nxt
            index += 1
        paragraph = doc.add_paragraph()
        if re.match(r"^\[\d+\]", combined):
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.0
            _add_inline(paragraph, combined, size=8.5 if compact else 9)
        else:
            _add_inline(paragraph, combined)


def _add_title_block(
    doc: Document,
    *,
    kicker: str,
    title: str,
    subtitle: str,
    authors: str,
    compact: bool,
    minimal: bool = False,
    include_date: bool = True,
) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2 if minimal else 14 if compact else 8)
    kicker_p = doc.add_paragraph()
    kicker_p.paragraph_format.space_after = Pt(5)
    kicker_run = kicker_p.add_run(kicker.upper())
    _set_run_font(kicker_run, size=9, color=CYAN, bold=True)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(7)
    title_p.paragraph_format.keep_with_next = True
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run(title)
    _set_run_font(title_run, size=18 if minimal else 22, color=NAVY, bold=True)

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(12)
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle_run = subtitle_p.add_run(subtitle)
    _set_run_font(subtitle_run, size=10.5 if minimal else 12.5, color=MUTED, italic=True)

    if authors:
        author_p = doc.add_paragraph()
        author_p.paragraph_format.space_after = Pt(3)
        author_run = author_p.add_run(authors)
        _set_run_font(author_run, size=9 if minimal else 10.5, color=INK, bold=True)
    if include_date:
        date_run = doc.add_paragraph().add_run("August 2026")
        _set_run_font(date_run, size=8.5 if minimal else 9.5, color=MUTED)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(6)
    rule.paragraph_format.space_after = Pt(4)
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), CYAN)
    p_bdr.append(bottom)
    rule._p.get_or_add_pPr().append(p_bdr)


def _set_core_properties(doc: Document, *, title: str, subject: str) -> None:
    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = "Andrei Pomorov"
    props.keywords = "distributed order management, MILP, LNS, QAOA, WISER"
    props.comments = "Final submission artifact generated from reviewed aggregate evidence."


def build_report() -> Path:
    doc = Document()
    _setup_styles(doc, compact=True)
    _configure_section(
        doc.sections[0],
        label="WISER DOM | Final Research Report",
        subtitle="Nestlé WISER Quantum Challenge",
        first_page=True,
    )
    _set_core_properties(
        doc,
        title="A Scalable Safeguarded Hybrid Classical-Quantum Solver for Distributed Order Management",
        subject="Final research report",
    )
    _add_title_block(
        doc,
        kicker="Final research report",
        title="A Scalable Safeguarded Hybrid Classical-Quantum Solver for Distributed Order Management",
        subtitle="Exact recourse, adaptive neighborhood search, and bounded quantum proposals",
        authors="Andrei Pomorov",
        compact=True,
    )
    _parse_markdown(
        doc,
        REPORTS / "final_report.md",
        compact=True,
        front_matter_until="Abstract",
    )
    output = REPORTS / "final_report.docx"
    doc.save(output)
    return output


def build_challenge_report() -> Path:
    doc = Document()
    _setup_styles(doc, compact=True)
    _configure_section(
        doc.sections[0],
        label="WISER DOM | Challenge Submission Report",
        subtitle="Requirement-by-requirement technical evidence",
        first_page=True,
    )
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(1.1)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    _set_core_properties(
        doc,
        title="WISER Quantum Challenge Submission Report",
        subject="Six-to-ten-page challenge criteria report",
    )
    _add_title_block(
        doc,
        kicker="Challenge submission report",
        title="Scalable, Safeguarded DOM Optimization",
        subtitle="Business value, mathematical model, implementation, scaling, and hardware evidence",
        authors="Andrei Pomorov",
        compact=True,
        minimal=True,
    )
    _parse_markdown(
        doc,
        REPORTS / "challenge_submission_report.md",
        compact=True,
        heading_shift=-1,
        front_matter_until="Portal summary",
        forced_break_before={
            "6.2 Coordinated improvement control",
            "7.4 IBM hardware and runtime correction",
            "8. Recommendation, limitations, and submission map",
        },
        figure_width_inches=5.1,
    )
    output = REPORTS / "challenge_submission_report.docx"
    doc.save(output)
    return output


def build_summary() -> Path:
    doc = Document()
    _setup_styles(doc, compact=True)
    _configure_section(
        doc.sections[0],
        label="WISER DOM | Business + Technical Summary",
        subtitle="Final submission | Andrei Pomorov",
        first_page=False,
    )
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    _set_core_properties(doc, title="WISER DOM Hybrid Solver", subject="Two-page business and technical summary")
    _add_title_block(
        doc,
        kicker="Business + technical summary",
        title="WISER DOM Hybrid Solver",
        subtitle="A portable classical default with safeguarded quantum experimentation",
        authors="",
        compact=True,
        minimal=True,
        include_date=False,
    )
    _parse_markdown(
        doc,
        REPORTS / "business_technical_summary.md",
        compact=True,
        heading_shift=-1,
    )
    output = REPORTS / "business_technical_summary.docx"
    doc.save(output)
    return output


def build_planner() -> Path:
    doc = Document()
    _setup_styles(doc, compact=True)
    normal = doc.styles["Normal"]
    normal.font.size = Pt(8.5)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.0
    for name, size, before, after in (
        ("Heading 1", 13, 8, 4),
        ("Heading 2", 11.5, 6, 3),
        ("Heading 3", 10.5, 4, 2),
    ):
        style = doc.styles[name]
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    _configure_section(
        section,
        label="WISER DOM | Planner Decision View",
        subtitle="Reviewed 100-group scope",
        first_page=False,
    )
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    _set_core_properties(doc, title="WISER DOM Planner Decision View", subject="One-page planner view")
    _add_title_block(
        doc,
        kicker="Planner decision view",
        title="Release the polished-greedy plan",
        subtitle="Exact recourse completed · independent validation passed · quality escalation available",
        authors="Andrei Pomorov | Reviewed scope: 100 assignment groups",
        compact=True,
        minimal=True,
        include_date=False,
    )
    _parse_markdown(doc, REPORTS / "planner_view.md", compact=True, heading_shift=-1)
    output = REPORTS / "planner_view.docx"
    doc.save(output)
    return output


def main() -> int:
    for builder in (build_report, build_challenge_report, build_summary, build_planner):
        print(builder().relative_to(ROOT))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
